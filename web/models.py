import datetime
from django.contrib.auth.models import User
from django.db import models
from django.urls import reverse
from django.utils.html import mark_safe
from django.conf import settings
from .LabelDecoder import decode_label_detail
from docx import Document
from docx.shared import Inches


class Diagnosis(models.Model):
    name = models.CharField(max_length=200, null=False)
    description = models.TextField(null=True, blank=True)
    text_to_onclusion = models.TextField(null=True, blank=True)

    def __str__(self):
        return self.name


class TemporaryContraindications(models.Model):
    name = models.CharField(max_length=200, null=False)
    description = models.TextField(null=True, blank=True)
    text_to_onclusion = models.TextField(null=True, blank=True)

    def __str__(self):
        return self.name


class Patient(models.Model):
    full_name = models.CharField(max_length=100, null=True, blank=True)
    age = models.PositiveSmallIntegerField(null=True, blank=True)

    diagnoses = models.ManyToManyField(Diagnosis, blank=True)

    def __str__(self):
        return self.full_name

    def get_diagnoses_name_list(self):
        if self.diagnoses:
            return [i.name for i in self.diagnoses.all()]
        return []


class Doctor(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)

    full_name = models.CharField(max_length=100, null=False, blank=False, default='Фамилия имя отчество')
    qualification = models.CharField(max_length=100, null=True, blank=True)
    experience = models.PositiveSmallIntegerField(null=True, blank=True)
    work_place = models.CharField(max_length=100, null=True, blank=True)
    education = models.CharField(max_length=100, null=True, blank=True)
    contacts = models.CharField(max_length=100, null=True, blank=True)
    photo = models.ImageField(null=True, default='profile_photo.jpg')

    def __str__(self):
        return f'{self.full_name} {self.user.username}'


class RecommendText(models.Model):
    operation = models.CharField(max_length=100)
    treatment_tactics = models.TextField()
    tactics_if_agree = models.TextField()


class NeuronetPrediction(models.Model):
    confidence = models.FloatField(null=True)

    recommend_text = models.OneToOneField(RecommendText, on_delete=models.CASCADE, null=True)

    def __str__(self):
        return f'{decode_label_detail(self.classification_types)} : {format(self.confidence, ".2f")}%'


class ClassificationType(models.Model):
    value = models.PositiveSmallIntegerField()
    prediction = models.ForeignKey(NeuronetPrediction, on_delete=models.CASCADE,
                                   related_name='classification_types')


def rename_file_by_pk(instance, filename):
    ext = filename.split('.')[-1]
    return 'snapshots/{}.{}'.format(Treatment.objects.count() + 1, ext)


class Treatment(models.Model):
    neurological_deficit = models.PositiveSmallIntegerField(null=False, default=1)
    conscious_level = models.PositiveSmallIntegerField(null=False, default=15)

    time_passed = models.PositiveSmallIntegerField(null=True, blank=True)
    hematoma_volume = models.PositiveSmallIntegerField(null=True, blank=True)

    is_injury = models.BooleanField(default=False, null=True, blank=True)
    has_stroke_symptoms = models.BooleanField(default=False, null=False)

    temporary_contraindications = models.ManyToManyField(TemporaryContraindications, blank=True)

    snapshot = models.ImageField(upload_to=rename_file_by_pk, null=True, blank=True)
    patient = models.ForeignKey(Patient, on_delete=models.CASCADE, null=True, blank=True)
    doctor = models.ForeignKey(Doctor, on_delete=models.CASCADE, null=True, default=None, blank=True)

    predict = models.OneToOneField(NeuronetPrediction,
                                   on_delete=models.CASCADE, null=True, blank=True)

    def __str__(self):
        patology = ''
        if self.predict:
            patology = decode_label_detail(ClassificationType.objects.filter(prediction=self.predict).all())
        return f'{self.patient.full_name}.{patology} '

    def get_snapshot_html(self):
        return mark_safe(f'<img src="{settings.MEDIA_URL}{self.snapshot}" alt="File image"/>')

    def get_snapshot_for_admin(self):
        HEIGHT = 120
        WIDTH = 200
        return mark_safe(f'<img src="{settings.MEDIA_URL}{self.snapshot}" '
                         f'style=width:{WIDTH}px; height:{HEIGHT}px; alt="File image"/>')

    def get_absolute_url(self):
        return reverse('treatment_detail', args=[str(self.id)])

    def get_temporary_contraindications_name_list(self):
        if self.temporary_contraindications:
            return [i.name for i in self.temporary_contraindications.all()]
        return []

    def formalize_to_document(self, file_name) -> str:
        self.document = Document()
        self.document.add_heading('Заключение', 0)

        p1 = self.document.add_paragraph('')
        p1.add_run("Пациент: ").bold = True
        p1.add_run(self.patient.full_name)

        p2 = self.document.add_paragraph('')
        p2.add_run("Дата рождения: ").bold = True

        p3 = self.document.add_paragraph('')
        p3.add_run("Возраст: ").bold = True
        p3.add_run(str(self.patient.age) + " лет")

        p4 = self.document.add_paragraph('')
        p4.add_run("Дата консультации: ").bold = True
        now = datetime.datetime.now()
        self.date = str(now.strftime("%d-%m-%Y %H:%M")).split()[0]
        p4.add_run(self.date)

        p5 = self.document.add_paragraph('')
        p5.add_run("Анамнез: ").bold = True

        if self.patient.diagnoses.all():
            self.document.add_paragraph("Сопутствующие патологии: ")
            for i in self.patient.diagnoses.all():
                self.document.add_paragraph(i.name, style='List Bullet')

        if self.temporary_contraindications.all():
            self.document.add_paragraph("Противопоказания: ")
            for i in self.temporary_contraindications.all():
                self.document.add_paragraph(i, style='List Bullet')

        p6 = self.document.add_paragraph('')
        p6.add_run("Данные обследования: ").bold = True
        self.document.add_paragraph("Уровень сознания: " + str(self.conscious_level))
        self.document.add_paragraph("Неврологический дефицит: " + str(self.neurological_deficit))
        if self.has_stroke_symptoms:
            self.document.add_paragraph("Присутствуют симптомы инсульта")
        else:
            self.document.add_paragraph("Симптомы инсульта отсутствуют")
        if self.is_injury:
            self.document.add_paragraph("Была получена травма")

        p7 = self.document.add_paragraph('')
        p7.add_run("Диагноз: ").bold = True
        p7.add_run(decode_label_detail(ClassificationType.objects.filter(prediction=self.predict).all()))

        p8 = self.document.add_paragraph('')
        p8.add_run("Рекомендовано: ").bold = True

        if self.predict.recommend_text.treatment_tactics:
            for recommend in self.predict.recommend_text.treatment_tactics.split('.')[:-1]:
                self.document.add_paragraph(recommend + '.', style='List Bullet')

        path = f'documents/{file_name}.docx'
        self.document.save(path)
        return path


class FAQ(models.Model):
    name = models.CharField(max_length=100, null=True, blank=True)

    def __str__(self):
        return self.name


class FAQItem(models.Model):
    header = models.CharField(max_length=100, null=False)
    text = models.TextField(null=False)
    reference = models.ForeignKey(FAQ, on_delete=models.CASCADE)

    def __str__(self):
        return self.header
